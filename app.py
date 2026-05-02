import io
import re
from pathlib import Path

import google.generativeai as genai
import streamlit as st
from docx import Document
from PIL import Image


def _split_table_row(row: str) -> list[str]:
    row = row.strip()
    if row.startswith("|"):
        row = row[1:]
    if row.endswith("|"):
        row = row[:-1]
    return [c.strip() for c in row.split("|")]


def _is_markdown_table_separator(line: str) -> bool:
    if "|" not in line:
        return False
    cells = _split_table_row(line)
    if not cells:
        return False
    for c in cells:
        t = c.strip().replace(" ", "")
        if not re.match(r"^:?-{2,}:?$", t):
            return False
    return True


def _parse_markdown_segments(markdown: str) -> list[tuple[str, object]]:
    lines = markdown.splitlines()
    segments: list[tuple[str, object]] = []
    i = 0
    while i < len(lines):
        if not lines[i].strip():
            i += 1
            continue

        stripped = lines[i].strip()
        if (
            "|" in stripped
            and i + 1 < len(lines)
            and _is_markdown_table_separator(lines[i + 1])
        ):
            header = _split_table_row(lines[i])
            i += 2
            rows: list[list[str]] = [header]
            while i < len(lines):
                s = lines[i].strip()
                if not s:
                    break
                if "|" not in s:
                    break
                if _is_markdown_table_separator(lines[i]):
                    i += 1
                    continue
                rows.append(_split_table_row(lines[i]))
                i += 1
            segments.append(("table", rows))
            continue

        block_lines: list[str] = []
        while i < len(lines):
            s = lines[i]
            if not s.strip():
                i += 1
                break
            stp = s.strip()
            if "|" in stp and i + 1 < len(lines) and _is_markdown_table_separator(
                lines[i + 1]
            ):
                break
            block_lines.append(s)
            i += 1
        if block_lines:
            segments.append(("text", "\n".join(block_lines)))

    return segments


def markdown_to_docx_buffer(markdown: str) -> io.BytesIO:
    doc = Document()
    for kind, payload in _parse_markdown_segments(markdown):
        if kind == "text":
            block = str(payload)
            current_para: list[str] = []

            def flush_para() -> None:
                if current_para:
                    doc.add_paragraph("\n".join(current_para))
                    current_para.clear()

            for line in block.split("\n"):
                s = line.strip()
                if not s:
                    flush_para()
                    continue
                if s.startswith("#"):
                    flush_para()
                    level = len(s) - len(s.lstrip("#"))
                    level = min(max(level, 1), 3)
                    doc.add_heading(s.lstrip("#").strip(), level=level)
                else:
                    current_para.append(line.rstrip())
            flush_para()

        elif kind == "table":
            rows: list[list[str]] = payload  # type: ignore[assignment]
            if not rows:
                continue
            num_cols = max(len(r) for r in rows)
            norm_rows: list[list[str]] = []
            for r in rows:
                extended = list(r) + [""] * (num_cols - len(r))
                norm_rows.append(extended[:num_cols])

            table = doc.add_table(rows=len(norm_rows), cols=num_cols)
            table.style = "Table Grid"

            for ri, row_cells in enumerate(norm_rows):
                for ci, cell_text in enumerate(row_cells):
                    cell = table.rows[ri].cells[ci]
                    cell.text = cell_text
                    if ri == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

            doc.add_paragraph()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


st.set_page_config(page_title="시험지 깔끔 변환기", layout="centered")

GEMINI_PROMPT = (
    "이 이미지들에서 볼펜이나 연필로 쓰인 낙서, 빗금, 풀이 과정, 체크 표시, 필기 흔적은 모두 무시해. "
    "단원명, 교재명, 페이지 번호, 머리말, 꼬리말, 학습 목표, 안내 문구처럼 문제 풀이에 직접 필요하지 않은 정보는 추출하지 마. "
    "문제번호, 문제 본문, 보기(①, ②, ③...), 조건, 지문, 그리고 표만 추출해. "
    "문제 본문과 보기는 마크다운(Markdown) 문법을 사용하지 말고 일반 텍스트(plain text)로 출력해. "
    "표가 있는 경우에만 마크다운(Markdown)의 Table 형식으로 정확하게 구현해. "
    "표 안에서 원화 금액처럼 숫자 앞에 ₩ 또는 \\ 기호가 붙은 값은 금액으로 판단하고, 해당 금액 셀은 오른쪽 정렬되도록 마크다운 표 정렬 구문을 사용해. "
    "예를 들어 금액 열의 구분선은 |---:| 형태로 작성해. "
    "사진은 업로드된 순서대로 이어진 페이지라고 간주해. "
    "한 문제가 여러 사진이나 여러 페이지에 걸쳐 이어져 있으면, 같은 문제번호의 내용으로 판단하여 하나의 문제로 합쳐서 출력해. "
    "다음 사진에 이어지는 문제의 나머지 내용, 보기, 표가 있으면 이전 사진의 해당 문제 아래에 이어 붙여. "
    "새 문제번호가 나오기 전까지는 같은 문제의 연속 내용으로 처리해. "
    "출력에는 추출한 문제 내용만 포함하고, 설명이나 요약, 추출 과정에 대한 말은 쓰지 마."
)


with st.sidebar:
    api_key = st.text_input(
        "Gemini API Key",
        type="password",
        help="Google AI Studio에서 발급받은 API 키를 입력하세요.",
    )

st.title("시험지 깔끔 변환기")

uploaded_file = st.file_uploader(
    "JPG, PNG, PDF 파일을 선택하세요",
    type=["jpg", "jpeg", "png", "pdf"],
)

key_ok = bool(api_key and api_key.strip())
file_ok = uploaded_file is not None
can_convert = key_ok and file_ok

if st.button("변환 시작", disabled=not can_convert):
    st.session_state.pop("gemini_markdown", None)
    file_bytes = uploaded_file.read()
    uploaded_stem = Path(uploaded_file.name or "시험지").stem
    safe_stem = re.sub(r'[\\/:*?"<>|]+', "_", uploaded_stem).strip() or "시험지"
    st.session_state["download_file_name"] = f"{safe_stem}_변환.docx"
    name_lower = (uploaded_file.name or "").lower()
    is_pdf = uploaded_file.type == "application/pdf" or name_lower.endswith(".pdf")

    genai.configure(api_key=api_key.strip())
    model = genai.GenerativeModel("gemini-2.5-flash")

    with st.spinner("변환 중..."):
        try:
            if is_pdf:
                content = [
                    GEMINI_PROMPT,
                    {"mime_type": "application/pdf", "data": file_bytes},
                ]
            else:
                image = Image.open(io.BytesIO(file_bytes))
                content = [GEMINI_PROMPT, image]

            response = model.generate_content(content)
            try:
                result_text = response.text
            except ValueError:
                result_text = None
                st.error("응답을 가져올 수 없습니다. 안전 필터 또는 빈 응답일 수 있습니다.")

            if result_text:
                st.session_state["gemini_markdown"] = result_text
        except Exception as e:
            st.error(f"API 호출 오류: {e}")

if st.session_state.get("gemini_markdown"):
    st.divider()
    st.subheader("추출 결과")
    st.markdown(st.session_state["gemini_markdown"])

    doc_buffer = markdown_to_docx_buffer(st.session_state["gemini_markdown"])
    st.download_button(
        label="워드 파일 다운로드",
        data=doc_buffer,
        file_name=st.session_state.get("download_file_name", "시험지_변환.docx"),
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
